from io import BytesIO
import unittest

from docx import Document as WordDocument

from csv_reviewer.document_generator import (
    build_docx,
    document_text,
    generate_document,
    update_section,
)


class DocumentGeneratorTests(unittest.TestCase):
    def setUp(self) -> None:
        self.metadata = {
            "document_type": "Validation Plan",
            "title": "NovaQMS Validation Plan",
            "document_id": "VP-001",
            "revision": "0.1",
            "system": "NovaQMS 2.4",
            "owner": "CSV Lead",
            "author": "Qualified Author",
            "purpose": "Define the validation approach for NovaQMS.",
            "scope": "The regulated quality workflow and its controlled interfaces.",
            "risk_classification": "GxP / High",
        }
        self.sources = [
            {
                "name": "VP-Template.docx",
                "role": "Template",
                "revision": "3.1",
                "text": "1. Purpose\n2. Scope\n3. Risk Management\n4. Acceptance Criteria",
            },
            {
                "name": "CSV-SOP.txt",
                "role": "Procedure",
                "revision": "5.2",
                "text": (
                    "The validation owner shall maintain traceability from approved requirements and risks to executed tests. "
                    "Acceptance criteria require approved deliverables, passed tests, and closed critical deviations."
                ),
            },
            {
                "name": "Project-Plan.txt",
                "role": "Plan",
                "revision": "1.0",
                "text": "NovaQMS is a high-risk GxP system. The scope includes SSO, audit trails, and electronic signatures.",
            },
        ]

    def test_template_structure_and_section_citations_are_retained(self) -> None:
        draft = generate_document(self.metadata, self.sources)

        self.assertEqual([section["title"] for section in draft["sections"]], ["Purpose", "Scope", "Risk Management", "Acceptance Criteria"])
        self.assertEqual(draft["generation_provenance"]["structure_source"], "SRC-001")
        self.assertIn("SRC-", " ".join(document_text(draft).split()))
        self.assertTrue(draft["generation_provenance"]["human_review_required"])
        self.assertNotIn("text", draft["sources"][0])

    def test_human_review_state_requires_every_section(self) -> None:
        draft = generate_document(self.metadata, self.sources)
        for section in draft["sections"]:
            draft = update_section(draft, section["section_id"], section["content"], "Quality Reviewer", reviewed=True)

        self.assertEqual(draft["status"], "Ready for Independent Review")
        self.assertTrue(all(section["status"] == "Human Reviewed" for section in draft["sections"]))
        self.assertEqual(len(draft["revision_history"]), len(draft["sections"]) + 1)

    def test_word_export_contains_controls_sections_and_provenance(self) -> None:
        draft = generate_document(self.metadata, self.sources)
        exported = build_docx(draft)
        document = WordDocument(BytesIO(exported))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertGreater(len(exported), 10_000)
        self.assertIn("NovaQMS Validation Plan", text)
        self.assertIn("DRAFT - HUMAN REVIEW REQUIRED", text)
        self.assertIn("1. Purpose", text)
        self.assertIn("Source and Generation Provenance", text)
        self.assertGreaterEqual(len(document.tables), 1)


if __name__ == "__main__":
    unittest.main()
