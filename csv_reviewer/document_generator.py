"""Source-grounded controlled document generation and Word export.

The generator is intentionally deterministic. It extracts structure from an
uploaded template, retrieves relevant statements from procedures/plans, and
creates an editable draft with section-level provenance. It does not approve,
publish, or silently invent controlled requirements.
"""

from __future__ import annotations

from copy import deepcopy
from datetime import UTC, datetime
import io
import re
from typing import Any
from uuid import uuid4

from docx import Document as WordDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCUMENT_BLUEPRINTS: dict[str, tuple[str, ...]] = {
    "Validation Plan": (
        "Purpose", "Scope", "System Overview", "Roles and Responsibilities",
        "Validation Strategy", "Deliverables", "Risk Management", "Traceability",
        "Schedule and Milestones", "Deviation Management", "Acceptance Criteria",
        "Approval and Release",
    ),
    "Validation Protocol": (
        "Purpose", "Scope", "Responsibilities", "Prerequisites", "Test Environment",
        "Test Procedures", "Expected Results", "Deviation Handling", "Acceptance Criteria",
        "Attachments and Evidence", "Approval and Release",
    ),
    "Validation Summary Report": (
        "Purpose", "Scope", "Validation Activities Performed", "Requirements Coverage",
        "Risk Coverage", "Deviations and Exceptions", "Results Summary", "Residual Risk",
        "Conclusion", "Approval and Release",
    ),
    "Standard Operating Procedure": (
        "Purpose", "Scope", "Definitions", "Roles and Responsibilities", "Procedure",
        "Exceptions and Escalation", "Records and Retention", "Training", "References",
        "Revision History",
    ),
    "Work Instruction": (
        "Purpose", "Scope", "Prerequisites", "Responsibilities", "Step-by-Step Instructions",
        "Expected Results", "Troubleshooting and Escalation", "Records", "References",
    ),
    "Risk Assessment": (
        "Purpose", "Scope", "Methodology", "Risk Criteria", "Hazards and Failure Modes",
        "Risk Controls", "Verification of Controls", "Residual Risk", "Approval and Release",
    ),
    "Test Plan": (
        "Purpose", "Scope", "Test Objectives", "Test Strategy", "Test Environment",
        "Roles and Responsibilities", "Test Data", "Entry Criteria", "Exit Criteria",
        "Defect Management", "Schedule", "Deliverables", "Approval and Release",
    ),
    "Test Protocol": (
        "Purpose", "Scope", "Prerequisites", "Test Data", "Test Steps", "Expected Results",
        "Evidence Requirements", "Deviation Handling", "Acceptance Criteria", "Approval and Release",
    ),
    "Requirements Specification": (
        "Purpose", "Scope", "System Overview", "Assumptions and Constraints",
        "Functional Requirements", "Data and Record Requirements", "Security Requirements",
        "Interface Requirements", "Performance Requirements", "Traceability", "Approval and Release",
    ),
    "Change Control Plan": (
        "Purpose", "Scope", "Change Description", "Business Justification", "Impact Assessment",
        "Risk Assessment", "Implementation Plan", "Verification Plan", "Rollback Plan",
        "Training and Communication", "Approval and Release",
    ),
}

SOURCE_ROLES = ("Template", "Procedure", "Plan", "Reference")
MAX_SECTIONS = 18
MAX_SOURCE_SENTENCES = 4


SECTION_KEYWORDS: dict[str, tuple[str, ...]] = {
    "purpose": ("purpose", "objective", "intent", "goal"),
    "scope": ("scope", "boundary", "included", "excluded", "applicable"),
    "roles": ("role", "responsibility", "owner", "author", "reviewer", "approver"),
    "responsibilities": ("role", "responsibility", "owner", "author", "reviewer", "approver"),
    "risk": ("risk", "hazard", "impact", "severity", "mitigation", "control"),
    "traceability": ("traceability", "requirement", "risk", "test", "coverage"),
    "test": ("test", "verify", "verification", "expected", "evidence", "result"),
    "acceptance": ("acceptance", "criteria", "pass", "fail", "approval"),
    "deviation": ("deviation", "exception", "defect", "issue", "capa"),
    "records": ("record", "retain", "retention", "evidence", "archive"),
    "training": ("training", "qualified", "competency"),
    "security": ("security", "access", "authentication", "authorization", "audit"),
    "schedule": ("schedule", "milestone", "timeline", "date", "phase"),
}


def _now() -> str:
    return datetime.now(UTC).isoformat()


def _display_timestamp(value: str) -> str:
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).strftime("%Y-%m-%d %H:%M UTC")
    except ValueError:
        return value


def _clean(value: Any, limit: int = 8_000) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()[:limit]


def _tokens(value: str) -> set[str]:
    stop = {"and", "the", "with", "from", "this", "that", "shall", "document", "section"}
    return {item for item in re.findall(r"[a-z0-9]{3,}", value.lower()) if item not in stop}


def _sentences(text: str) -> list[str]:
    output: list[str] = []
    for block in re.split(r"[\r\n]+", text or ""):
        for sentence in re.split(r"(?<=[.!?])\s+", block):
            cleaned = _clean(sentence, 700)
            if 24 <= len(cleaned) <= 700 and cleaned not in output:
                output.append(cleaned)
    return output


def extract_template_sections(template_text: str) -> list[str]:
    """Extract likely headings without treating template content as instructions."""

    headings: list[str] = []
    for raw in (template_text or "").splitlines():
        line = re.sub(r"^\s*(?:section\s+)?\d+(?:\.\d+)*[.):]?\s*", "", raw, flags=re.I).strip()
        line = re.sub(r"\s+", " ", line)
        if not 3 <= len(line) <= 80 or len(line.split()) > 10:
            continue
        if line.endswith((".", ";", ",")):
            continue
        alpha = sum(character.isalpha() for character in line)
        if alpha < 3:
            continue
        looks_like_heading = raw.strip().isupper() or raw.strip()[:1].isdigit() or line.istitle()
        if looks_like_heading and line.lower() not in {item.lower() for item in headings}:
            headings.append(line)
        if len(headings) >= MAX_SECTIONS:
            break
    return headings if len(headings) >= 3 else []


def _source_catalog(sources: list[dict[str, Any]]) -> list[dict[str, Any]]:
    catalog = []
    for index, source in enumerate(sources, start=1):
        role = source.get("role") if source.get("role") in SOURCE_ROLES else "Reference"
        catalog.append(
            {
                "source_id": f"SRC-{index:03d}",
                "name": _clean(source.get("name") or f"Source {index}", 240),
                "role": role,
                "revision": _clean(source.get("revision") or "Not identified", 100),
                "text": str(source.get("text") or "")[:500_000],
                "extraction_status": _clean(source.get("extraction_status") or "Complete", 40),
                "warnings": list(source.get("warnings") or []),
            }
        )
    return catalog


def _section_evidence(title: str, sources: list[dict[str, Any]]) -> list[dict[str, str]]:
    title_tokens = _tokens(title)
    keywords = set(title_tokens)
    lowered = title.lower()
    for key, values in SECTION_KEYWORDS.items():
        if key in lowered:
            keywords.update(values)
    ranked: list[tuple[float, str, str]] = []
    for source in sources:
        role_weight = {"Procedure": 0.35, "Plan": 0.25, "Reference": 0.15, "Template": 0.05}[source["role"]]
        for position, sentence in enumerate(_sentences(source["text"])):
            sentence_tokens = _tokens(sentence)
            overlap = len(keywords & sentence_tokens)
            if not overlap:
                continue
            score = overlap / max(1, len(keywords)) + role_weight - min(position, 80) / 1_000
            ranked.append((score, source["source_id"], sentence))
    ranked.sort(key=lambda item: item[0], reverse=True)
    selected: list[dict[str, str]] = []
    seen: set[str] = set()
    for _, source_id, sentence in ranked:
        fingerprint = sentence.lower()
        if fingerprint in seen:
            continue
        selected.append({"source_id": source_id, "text": sentence})
        seen.add(fingerprint)
        if len(selected) >= MAX_SOURCE_SENTENCES:
            break
    return selected


def _section_content(title: str, metadata: dict[str, str], evidence: list[dict[str, str]], sources: list[dict[str, Any]]) -> str:
    lower = title.lower()
    if "purpose" in lower and metadata.get("purpose"):
        lead = metadata["purpose"]
    elif "scope" in lower and metadata.get("scope"):
        lead = metadata["scope"]
    elif "approval" in lower or "release" in lower:
        lead = (
            "This draft requires review and approval through the organization’s controlled quality process. "
            "CSVQualReviewer does not approve, publish, or release this document."
        )
    elif "responsibilit" in lower or "roles" in lower:
        lead = (
            f"The document owner is {metadata.get('owner') or 'to be assigned'}. "
            "Authors, reviewers, approvers, and activity owners shall be identified before execution."
        )
    else:
        lead = (
            f"This {title.lower()} section was assembled from the controlled source set below. "
            "The author shall confirm applicability, completeness, and organization-specific terminology."
        )

    paragraphs = [lead]
    if evidence:
        paragraphs.append("Grounded source statements:")
        paragraphs.extend(f"- {item['text']} [{item['source_id']}]" for item in evidence)
    else:
        role_names = [source["name"] for source in sources if source["role"] in {"Procedure", "Plan"}]
        if role_names:
            paragraphs.append(
                "No directly matching statement was identified automatically. Confirm this section against: "
                + "; ".join(role_names[:5])
                + "."
            )
        else:
            paragraphs.append("Author input required: no controlled procedure or plan was supplied for this topic.")
    return "\n\n".join(paragraphs)


def generate_document(metadata: dict[str, Any], sources: list[dict[str, Any]]) -> dict[str, Any]:
    title = _clean(metadata.get("title"), 240)
    document_type = _clean(metadata.get("document_type"), 120)
    if not title:
        raise ValueError("Enter a document title.")
    if document_type not in DOCUMENT_BLUEPRINTS:
        raise ValueError("Select a supported document type.")

    normalized_metadata = {
        "title": title,
        "document_type": document_type,
        "document_id": _clean(metadata.get("document_id") or "DRAFT", 80),
        "revision": _clean(metadata.get("revision") or "0.1", 40),
        "system": _clean(metadata.get("system") or "Not specified", 160),
        "owner": _clean(metadata.get("owner") or "To be assigned", 160),
        "author": _clean(metadata.get("author") or "Document Generation Agent", 160),
        "purpose": _clean(metadata.get("purpose"), 2_000),
        "scope": _clean(metadata.get("scope"), 2_000),
        "risk_classification": _clean(metadata.get("risk_classification") or "Not classified", 80),
    }
    catalog = _source_catalog(sources)
    template = next((source for source in catalog if source["role"] == "Template"), None)
    headings = extract_template_sections(template["text"] if template else "")
    structure_source = template["source_id"] if headings and template else "CSVQualReviewer controlled blueprint v1.0"
    if not headings:
        headings = list(DOCUMENT_BLUEPRINTS[document_type])

    sections = []
    for index, heading in enumerate(headings[:MAX_SECTIONS], start=1):
        evidence = _section_evidence(heading, catalog)
        sections.append(
            {
                "section_id": f"SEC-{index:03d}",
                "number": index,
                "title": heading,
                "content": _section_content(heading, normalized_metadata, evidence, catalog),
                "status": "Draft",
                "source_ids": sorted({item["source_id"] for item in evidence}),
                "generated_at": _now(),
                "reviewed_by": "",
                "reviewed_at": "",
            }
        )

    assumptions = []
    present_roles = {source["role"] for source in catalog}
    if "Template" not in present_roles:
        assumptions.append("No template was supplied; the controlled CSVQualReviewer blueprint was used.")
    if "Procedure" not in present_roles:
        assumptions.append("No approved procedure was supplied; procedural alignment requires author review.")
    if "Plan" not in present_roles:
        assumptions.append("No plan was supplied; project-specific scope, roles, and milestones require author review.")
    failed = [source["name"] for source in catalog if source["extraction_status"] == "Failed"]
    if failed:
        assumptions.append("Text extraction failed for: " + "; ".join(failed) + ".")

    created_at = _now()
    return {
        "schema_version": "1.0",
        "draft_id": f"DRAFT-{datetime.now(UTC).strftime('%Y%m%d-%H%M%S')}-{uuid4().hex[:6].upper()}",
        "status": "Draft - Human Review Required",
        "metadata": normalized_metadata,
        "sections": sections,
        "sources": [{key: value for key, value in source.items() if key != "text"} for source in catalog],
        "assumptions": assumptions,
        "generation_provenance": {
            "agent": "CSVQualReviewer Document Generation Agent",
            "mode": "Deterministic source-grounded drafting",
            "structure_source": structure_source,
            "created_at": created_at,
            "human_review_required": True,
            "approval_boundary": "Approval and release remain outside CSVQualReviewer.",
        },
        "revision_history": [
            {"event_id": f"GEN-{uuid4().hex[:10].upper()}", "timestamp": created_at, "user": normalized_metadata["author"], "action": "Draft generated", "detail": structure_source}
        ],
    }


def update_section(draft: dict[str, Any], section_id: str, content: str, reviewer: str, reviewed: bool = False) -> dict[str, Any]:
    updated = deepcopy(draft)
    section = next((item for item in updated["sections"] if item["section_id"] == section_id), None)
    if section is None:
        raise ValueError("The selected section was not found.")
    cleaned = str(content or "").strip()[:20_000]
    if not cleaned:
        raise ValueError("Section content cannot be empty.")
    timestamp = _now()
    section["content"] = cleaned
    section["status"] = "Human Reviewed" if reviewed else "Draft"
    section["reviewed_by"] = _clean(reviewer or "Not identified", 160) if reviewed else ""
    section["reviewed_at"] = timestamp if reviewed else ""
    updated["revision_history"].append(
        {
            "event_id": f"GEN-{uuid4().hex[:10].upper()}",
            "timestamp": timestamp,
            "user": _clean(reviewer or "Not identified", 160),
            "action": "Section reviewed" if reviewed else "Section edited",
            "detail": f"{section_id} · {section['title']}",
        }
    )
    reviewed_count = sum(item["status"] == "Human Reviewed" for item in updated["sections"])
    updated["status"] = (
        "Ready for Independent Review"
        if reviewed_count == len(updated["sections"])
        else "Draft - Human Review Required"
    )
    return updated


def document_text(draft: dict[str, Any]) -> str:
    metadata = draft["metadata"]
    lines = [
        metadata["title"],
        f"Document ID: {metadata['document_id']}",
        f"Revision: {metadata['revision']}",
        f"Status: {draft['status']}",
        "",
    ]
    for section in draft["sections"]:
        lines.extend([f"{section['number']}. {section['title']}", section["content"], ""])
    lines.append("Source provenance")
    lines.extend(f"{source['source_id']}: {source['name']} ({source['role']}, revision {source['revision']})" for source in draft["sources"])
    return "\n".join(lines)


def _set_font(run: Any, name: str, size: float, color: str = "17231E", bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell: Any, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    tc_width = properties.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        properties.append(tc_width)
    tc_width.set(qn("w:w"), str(width_dxa))
    tc_width.set(qn("w:type"), "dxa")


def _configure_table(table: Any, widths: tuple[int, ...]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            _set_cell_width(cell, value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_page_field(paragraph: Any, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _clear_template_body(document: Any) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def build_docx(draft: dict[str, Any], template_bytes: bytes | None = None) -> bytes:
    """Create a polished controlled-draft DOCX, preserving template styles when supplied."""

    used_template = False
    if template_bytes:
        try:
            document = WordDocument(io.BytesIO(template_bytes))
            _clear_template_body(document)
            used_template = True
        except Exception:
            document = WordDocument()
    else:
        document = WordDocument()

    section = document.sections[0]
    if not used_template:
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    styles = document.styles
    if not used_template:
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1
        for style_name, size, color, before, after in (
            ("Title", 24, "0B2545", 0, 8),
            ("Heading 1", 16, "2E74B5", 16, 8),
            ("Heading 2", 13, "2E74B5", 12, 6),
            ("Heading 3", 12, "1F4D78", 8, 4),
        ):
            style = styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "CSVQualReviewer | Controlled Draft"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        _set_font(run, "Calibri", 8.5, "68768A", True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = footer.add_run("Human review required  |  Page ")
    _set_font(prefix, "Calibri", 8.5, "68768A")
    _add_page_field(footer, "PAGE")
    separator = footer.add_run(" of ")
    _set_font(separator, "Calibri", 8.5, "68768A")
    _add_page_field(footer, "NUMPAGES")

    metadata = draft["metadata"]
    title = document.add_paragraph(style="Title")
    title_run = title.add_run(metadata["title"])
    if not used_template:
        _set_font(title_run, "Calibri", 24, "0B2545", True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run(f"{metadata['document_type']} | {draft['status']}")
    _set_font(subtitle_run, "Calibri", 12, "56657A", False)

    control = document.add_table(rows=0, cols=2)
    control.style = "Table Grid"
    fields = (
        ("Document ID", metadata["document_id"]), ("Revision", metadata["revision"]),
        ("System / Project", metadata["system"]), ("Risk Classification", metadata["risk_classification"]),
        ("Owner", metadata["owner"]), ("Author", metadata["author"]),
        ("Draft ID", draft["draft_id"]), ("Generated", _display_timestamp(draft["generation_provenance"]["created_at"])),
    )
    for label, value in fields:
        cells = control.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
        _set_cell_shading(cells[0], "E8EEF5")
        for run in cells[0].paragraphs[0].runs:
            _set_font(run, "Calibri", 9.5, "0B2545", True)
        for run in cells[1].paragraphs[0].runs:
            _set_font(run, "Calibri", 9.5, "17231E")
    _configure_table(control, (2700, 6660))

    notice = document.add_paragraph()
    notice.paragraph_format.space_before = Pt(12)
    notice.paragraph_format.space_after = Pt(10)
    notice_properties = notice._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF4CC")
    notice_properties.append(shading)
    notice_run = notice.add_run(
        "DRAFT - HUMAN REVIEW REQUIRED. Generated content must be verified against controlled sources before approval or use."
    )
    _set_font(notice_run, "Calibri", 10, "7A5A00", True)

    for item in draft["sections"]:
        document.add_heading(f"{item['number']}. {item['title']}", level=1)
        for block in item["content"].split("\n"):
            text = block.strip()
            if not text:
                continue
            if text.startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(text[2:])
            else:
                paragraph = document.add_paragraph(text)
            paragraph.paragraph_format.keep_together = False
        source_ids = ", ".join(item["source_ids"]) or "No direct source match - author confirmation required"
        citation = document.add_paragraph()
        citation.paragraph_format.space_before = Pt(4)
        citation.paragraph_format.space_after = Pt(4)
        citation_run = citation.add_run(f"Section provenance: {source_ids} | Status: {item['status']}")
        citation_run.italic = True
        _set_font(citation_run, "Calibri", 8.5, "68768A")

    document.add_heading("Source and Generation Provenance", level=1)
    document.add_paragraph(
        "The source list records the exact inputs available to the generation run. Source content remains authoritative; this draft is advisory until human review and approval."
    )
    for source in draft["sources"]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(
            f"{source['source_id']} - {source['name']} | {source['role']} | Revision {source['revision']} | Extraction {source['extraction_status']}"
        )
    if draft["assumptions"]:
        document.add_heading("Assumptions and Open Author Actions", level=2)
        for assumption in draft["assumptions"]:
            document.add_paragraph(assumption, style="List Bullet")
    document.add_heading("Generation Controls", level=2)
    for label, value in draft["generation_provenance"].items():
        if label == "created_at":
            value = _display_timestamp(str(value))
        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{label.replace('_', ' ').title()}: ")
        label_run.bold = True
        paragraph.add_run(str(value))

    document.core_properties.title = metadata["title"]
    document.core_properties.subject = f"{metadata['document_type']} controlled draft"
    document.core_properties.author = "CSVQualReviewer Document Generation Agent"
    document.core_properties.comments = "AI-assisted draft. Human review and controlled approval required."
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
